from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from aiogram.types import BufferedInputFile
from docx.shared import Mm
from docx import Document
from typing import List
import io


def create_word_document(
    buffered_files: list[BufferedInputFile],
    answers: list[list[str]] = None
) -> bytes:
    """
    Создает документ Word с изображениями (каждый вариант на отдельной альбомной странице A4 без полей)
    
    :param buffered_files: Список объектов BufferedInputFile с изображениями
    :param answers: Список списков с ответами для каждого варианта
    :return: Байты готового Word-документа
    """
    doc = Document()
    
    A4_WIDTH = Mm(210)
    A4_HEIGHT = Mm(297)
    
    for i, buffered_file in enumerate(buffered_files):
        try:
            image_data = buffered_file.data
        except AttributeError:
            image_data = buffered_file.read()

        if i == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section()
            
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = A4_WIDTH
        section.page_height = A4_HEIGHT

        section.left_margin = Mm(0)
        section.right_margin = Mm(0)
        section.top_margin = Mm(0)
        section.bottom_margin = Mm(0)
        section.header_distance = Mm(0)
        section.footer_distance = Mm(0)
        
        with io.BytesIO(image_data) as image_stream:
            doc.add_picture(image_stream, width=A4_WIDTH, height=A4_HEIGHT)
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        

    if answers:
        answers_section = doc.add_section()
        answers_section.orientation = WD_ORIENTATION.PORTRAIT
        answers_section.page_width = Mm(210)
        answers_section.page_height = Mm(297)
        answers_section.left_margin = Mm(15)
        answers_section.right_margin = Mm(15)
        answers_section.top_margin = Mm(15)
        answers_section.bottom_margin = Mm(15)
        
        doc.add_heading('Ответы к вариантам', level=1)

    table = doc.add_table(rows=len(answers[0]) + 1, cols=len(answers) + 1)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    for col_idx, variant_num in enumerate(range(1, len(answers) + 1), 1):
        hdr_cells[col_idx].text = f'Вар. {variant_num}'
        hdr_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for task_idx in range(len(answers[0])):
        row_cells = table.rows[task_idx + 1].cells
        row_cells[0].text = str(task_idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for variant_idx in range(len(answers)):
            cell = row_cells[variant_idx + 1]
            cell.text = answers[variant_idx][task_idx]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()


def create_two_vertical_A5_variants(
    buffered_files: List[BufferedInputFile],
    answers: List[List[str]] = None
) -> bytes:
    """
    Создает документ Word с изображениями (два варианта A5 на одном листе A4 в альбомной ориентации без полей)
    
    :param buffered_files: Список объектов BufferedInputFile с изображениями
    :param answers: Список списков с ответами для каждого варианта
    :return: Байты готового Word-документа
    """
    doc = Document()

    A4_WIDTH = Mm(297)
    A4_HEIGHT = Mm(210)
    A5_WIDTH = Mm(148.5)
    A5_HEIGHT = Mm(210)
    
    for i in range(0, len(buffered_files), 2):
        if i == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section()
            
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = A4_WIDTH
        section.page_height = A4_HEIGHT
        section.left_margin = Mm(0)
        section.right_margin = Mm(0)
        section.top_margin = Mm(0)
        section.bottom_margin = Mm(0)
        section.header_distance = Mm(0)
        section.footer_distance = Mm(0)
        
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        for cell in table.rows[0].cells:
            cell.width = A5_WIDTH
    
        try:
            image_data = buffered_files[i].data
        except AttributeError:
            image_data = buffered_files[i].read()
        
        with io.BytesIO(image_data) as image_stream:
            cell = table.rows[0].cells[0]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.paragraphs[0].add_run().add_picture(image_stream, width=A5_WIDTH, height=A5_HEIGHT)
        
        if i + 1 < len(buffered_files):
            try:
                image_data = buffered_files[i+1].data
            except AttributeError:
                image_data = buffered_files[i+1].read()
            
            with io.BytesIO(image_data) as image_stream:
                cell = table.rows[0].cells[1]
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell.paragraphs[0].add_run().add_picture(image_stream, width=A5_WIDTH, height=A5_HEIGHT)

    if answers:
        answers_section = doc.add_section()
        answers_section.orientation = WD_ORIENTATION.PORTRAIT
        answers_section.page_width = Mm(210)
        answers_section.page_height = Mm(297)
        answers_section.left_margin = Mm(15)
        answers_section.right_margin = Mm(15)
        answers_section.top_margin = Mm(15)
        answers_section.bottom_margin = Mm(15)
        
        doc.add_heading('Ответы к вариантам', level=1)

    table = doc.add_table(rows=len(answers[0]) + 1, cols=len(answers) + 1)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    for col_idx, variant_num in enumerate(range(1, len(answers) + 1), 1):
        hdr_cells[col_idx].text = f'Вар. {variant_num}'
        hdr_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    for task_idx in range(len(answers[0])):
        row_cells = table.rows[task_idx + 1].cells
        row_cells[0].text = str(task_idx + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for variant_idx in range(len(answers)):
            cell = row_cells[variant_idx + 1]
            cell.text = answers[variant_idx][task_idx]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes.getvalue()